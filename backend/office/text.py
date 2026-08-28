"""Yazı belgesi — gerçek .docx, gerekçeli değişikliklerle.

Tablo tarafıyla aynı sözleşme: gerçek dosya biçimi, doğrudan model erişimi,
her değişiklik neden yapıldığını taşıyor.

Paragraflar sıra numarasıyla adreslenir. Bu kırılgan görünüyor — araya
paragraf eklenince numaralar kayıyor — ama alternatifi her paragrafa kalıcı
kimlik basmak ve bunu .docx'e sığdırmak; taşınabilirliği bozardı. Bunun
yerine `read` her zaman güncel numaraları döndürüyor ve model düzenlemeden
önce okumak zorunda.
"""

from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path

import docx
from docx.document import Document as DocxDocument

from .model import OfficeDocument

MAX_PARAGRAPHS = 400


class TextError(RuntimeError):
    """Yazı belgesi işlemi yapılamadı."""


@dataclass
class TextDocument(OfficeDocument):
    doc: DocxDocument = field(default_factory=docx.Document)

    @property
    def kind(self) -> str:
        return "text"

    @classmethod
    def create(cls, path: str) -> TextDocument:
        return cls(path=path, doc=docx.Document())

    @classmethod
    def open(cls, path: str) -> TextDocument:
        target = Path(path)
        if not target.exists():
            raise TextError(f"{target} does not exist")
        try:
            return cls(path=str(target), doc=docx.Document(str(target)))
        except Exception as exc:
            raise TextError(f"could not open {target}: {exc}") from None

    # --- okuma ------------------------------------------------------------

    def summary(self) -> str:
        paragraphs = self.doc.paragraphs
        words = sum(len(p.text.split()) for p in paragraphs)
        lines = [
            f"{self.path} ({self.kind})",
            f"  {len(paragraphs)} paragraf, {len(self.doc.tables)} tablo, ~{words} kelime",
        ]
        if self.ledger.dirty:
            lines.append(f"  {self.ledger.unsaved_count} unsaved changes")
        return "\n".join(lines)

    def read(self, start: int = 0, count: int = MAX_PARAGRAPHS) -> str:
        paragraphs = self.doc.paragraphs
        if not paragraphs:
            return "(the document is empty)"

        end = min(start + count, len(paragraphs))
        lines = []
        for index in range(start, end):
            paragraph = paragraphs[index]
            style = paragraph.style.name if paragraph.style else "Normal"
            marker = f"[{index}]" if style == "Normal" else f"[{index}:{style}]"
            lines.append(f"{marker} {paragraph.text}")

        if end < len(paragraphs):
            lines.append(f"[... showing {end} of {len(paragraphs)} paragraphs]")
        return "\n".join(lines)

    # --- yazma ------------------------------------------------------------

    def append(self, text: str, why: str, style: str | None = None) -> str:
        try:
            paragraph = self.doc.add_paragraph(text, style=style)
        except KeyError:
            raise TextError(
                f"There is no style called {style!r}. The common ones are: Title, "
                f"Heading 1, Heading 2, Normal, List Bullet, Quote"
            ) from None
        index = len(self.doc.paragraphs) - 1
        self.ledger.record(f"paragraph {index}", None, text, why)
        return f"Paragraph {index} added ({paragraph.style.name}) — {why}"

    def replace(self, index: int, text: str, why: str) -> str:
        paragraphs = self.doc.paragraphs
        if not 0 <= index < len(paragraphs):
            raise TextError(
                f"There is no paragraph {index}; the document has {len(paragraphs)}"
            )
        paragraph = paragraphs[index]
        before = paragraph.text

        # Bir paragrafın metnini değiştirmenin tek yolu run'larını yeniden
        # kurmak. İlk run'a yazıp kalanları boşaltmak, paragrafın
        # biçimlendirmesini (yazı tipi, kalınlık) koruyor.
        if paragraph.runs:
            paragraph.runs[0].text = text
            for run in paragraph.runs[1:]:
                run.text = ""
        else:
            paragraph.add_run(text)

        self.ledger.record(f"paragraph {index}", before, text, why)
        return f"Paragraph {index} replaced — {why}"

    def add_table(self, rows: list[list[str]], why: str) -> str:
        if not rows or not rows[0]:
            raise TextError("A table needs at least one row and one column")
        table = self.doc.add_table(rows=len(rows), cols=len(rows[0]))
        table.style = "Table Grid"
        for r, row in enumerate(rows):
            for c, value in enumerate(row):
                if c < len(table.columns):
                    table.rows[r].cells[c].text = str(value)
        self.ledger.record(
            f"tablo {len(self.doc.tables) - 1}", None,
            f"{len(rows)}x{len(rows[0])}", why,
        )
        return f"{len(rows)}x{len(rows[0])} tablo eklendi — {why}"

    # --- kaydetme ve geri alma --------------------------------------------

    def save(self, path: str | None = None) -> str:
        target = Path(path or self.path)
        target.parent.mkdir(parents=True, exist_ok=True)
        try:
            self.doc.save(str(target))
        except Exception as exc:
            raise TextError(f"{target} kaydedilemedi: {exc}") from None
        self.path = str(target)
        self.ledger.mark_saved()
        return f"{target} saved ({len(self.ledger)} changes)."

    def undo(self, count: int = 1) -> str:
        """Yalnızca paragraf metni değişiklikleri geri alınabilir.

        Ekleme ve tablo geri alınmıyor: python-docx bir paragrafı gerçekten
        silmek için XML ağacına inmeyi gerektiriyor ve yarım çalışan bir geri
        alma, hiç çalışmayandan tehlikeli.
        """
        changes = self.ledger.last(count)
        if not changes:
            return "There is no change to revert."

        undone = 0
        for change in changes:
            if not change.target.startswith("paragraph ") or change.before is None:
                break
            index = int(change.target.split()[1])
            if index >= len(self.doc.paragraphs):
                break
            paragraph = self.doc.paragraphs[index]
            if paragraph.runs:
                paragraph.runs[0].text = change.before
                for run in paragraph.runs[1:]:
                    run.text = ""
            undone += 1

        self.ledger.drop_last(undone)
        if undone < len(changes):
            return (
                f"{undone} changes reverted. The rest were insertions, so they "
                f"could not be reverted — rebuild the document if you need them gone."
            )
        return f"{undone} changes reverted."
